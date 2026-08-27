from docx import Document

from resumeforge.loader import load_resume

from resumeforge.domain import (
    Certification,
    Project,
    Header,
    ResumeProfile,
)

from resumeforge.renderers.base import BaseRenderer
from resumeforge.renderers.header import HeaderRenderer
from resumeforge.renderers.summary import SummaryRenderer
from resumeforge.renderers.experience import ExperienceRenderer
from resumeforge.renderers.education import EducationRenderer
from resumeforge.renderers.skills import SkillsRenderer
from resumeforge.renderers.certification import CertificationRenderer
from resumeforge.renderers.project import ProjectRenderer
from resumeforge.layout import WordLayout


def test_header_renderer_is_base_renderer():
    assert issubclass(HeaderRenderer, BaseRenderer)

def test_summary_renderer_is_base_renderer():
    assert issubclass(SummaryRenderer, BaseRenderer)

def test_experience_renderer_is_base_renderer():
    assert issubclass(ExperienceRenderer, BaseRenderer)

def test_education_renderer_is_base_renderer():
    assert issubclass(EducationRenderer, BaseRenderer)

def test_skills_renderer_is_base_renderer():
    assert issubclass(SkillsRenderer, BaseRenderer)

def test_certification_renderer_is_base_renderer():
    assert issubclass(CertificationRenderer, BaseRenderer)

def test_project_renderer_is_base_renderer():
    assert issubclass(ProjectRenderer, BaseRenderer)

def test_header_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = load_resume()

    renderer = HeaderRenderer()

    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)

    assert "Jason Little" in text
    assert "Senior Software Engineer" in text

def test_summary_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = load_resume()

    renderer = SummaryRenderer()

    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)

    assert "Executive Summary" in text

def test_experience_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = load_resume()

    renderer = ExperienceRenderer()

    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)

    assert "Experience" in text
    assert "Regions Bank" in text
    assert "Senior Software Engineer" in text

def test_education_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = load_resume()

    renderer = EducationRenderer()

    renderer.render(layout, resume)

    assert "Education" in "\n".join(p.text for p in document.paragraphs)

def test_skills_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = load_resume()

    renderer = SkillsRenderer()

    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)

    assert "Technical Skills" in text
    assert "C#" in text

def test_certification_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = ResumeProfile(
        header=Header(
            name="Test",
            headline="Test",
            tagline="Test",
            location="",
            phone="",
            email="",
            linkedin="",
            github="",
            portfolio="",
        ),
        certifications=[
            Certification(
                name="Microsoft Certified: Azure Fundamentals",
                issuer="Microsoft",
                year="2024",
            )
        ],
    )

    renderer = CertificationRenderer()
    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)
    assert "Certifications" in text
    assert "Azure Fundamentals" in text

def test_project_renderer_runs():
    document = Document()
    layout = WordLayout(document)

    resume = ResumeProfile(
        header=Header(
            name="Test",
            headline="Test",
            tagline="Test",
            location="",
            phone="",
            email="",
            linkedin="",
            github="",
            portfolio="",
        ),
        projects=[
            Project(
                name="Test Project",
                description="A simple test project",
            )
        ],
    )

    renderer = ProjectRenderer()
    renderer.render(layout, resume)

    text = "\n".join(p.text for p in document.paragraphs)
    assert "Projects" in text
    assert "Test Project" in text