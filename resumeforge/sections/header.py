from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def render(document, resume):

    title = document.add_heading(resume["name"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(resume["headline"] + "\n")
    run.bold = True
    run.font.size = Pt(14)

    p.add_run(resume["tagline"])